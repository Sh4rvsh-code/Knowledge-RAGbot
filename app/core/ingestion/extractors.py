"""Text extraction from various document formats."""
from abc import ABC, abstractmethod
from typing import Dict, List, Any, Optional
import os
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
import chardet

from app.utils.logger import app_logger as logger


class ExtractionResult:
    """Structured output from text extraction."""
    
    def __init__(
        self,
        text: str,
        metadata: Dict[str, Any],
        page_mapping: Optional[List[Dict[str, Any]]] = None
    ):
        self.text = text
        self.metadata = metadata
        self.page_mapping = page_mapping or []
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "text": self.text,
            "metadata": self.metadata,
            "page_mapping": self.page_mapping
        }


class TextExtractor(ABC):
    """Abstract base class for text extraction."""
    
    @abstractmethod
    def extract(self, file_path: str) -> ExtractionResult:
        """Extract text from file."""
        pass
    
    @abstractmethod
    def supports(self, file_extension: str) -> bool:
        """Check if this extractor supports the file type."""
        pass


class PDFExtractor(TextExtractor):
    """Extract text from PDF files using PyMuPDF."""
    
    def supports(self, file_extension: str) -> bool:
        """Check if file is a PDF."""
        return file_extension.lower() == '.pdf'
    
    def extract(self, file_path: str) -> ExtractionResult:
        """
        Extract text from PDF page-by-page with metadata.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            ExtractionResult with text, metadata, and page mappings
        """
        logger.info(f"Extracting text from PDF: {file_path}")
        
        try:
            doc = fitz.open(file_path)
            all_text = []
            page_mapping = []
            current_char_position = 0
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                
                # Track character positions for highlighting
                page_start = current_char_position
                page_end = current_char_position + len(page_text)
                
                page_mapping.append({
                    "page_number": page_num + 1,
                    "start_char": page_start,
                    "end_char": page_end,
                    "text_length": len(page_text)
                })
                
                all_text.append(page_text)
                current_char_position = page_end + 1  # +1 for newline
            
            # Combine all pages
            full_text = "\n".join(all_text)
            
            # Extract metadata
            metadata = {
                "total_pages": len(doc),
                "file_size": os.path.getsize(file_path),
                "file_name": Path(file_path).name,
                "file_type": "pdf",
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "creator": doc.metadata.get("creator", ""),
            }
            
            doc.close()
            
            logger.info(f"Extracted {len(full_text)} characters from {metadata['total_pages']} pages")
            
            return ExtractionResult(
                text=full_text,
                metadata=metadata,
                page_mapping=page_mapping
            )
            
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract PDF: {str(e)}")


class DOCXExtractor(TextExtractor):
    """Extract text from DOCX files using python-docx."""
    
    def supports(self, file_extension: str) -> bool:
        """Check if file is a DOCX."""
        return file_extension.lower() in ['.docx', '.doc']
    
    def extract(self, file_path: str) -> ExtractionResult:
        """
        Extract text from DOCX including paragraphs and tables.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ExtractionResult with text and metadata
        """
        logger.info(f"Extracting text from DOCX: {file_path}")
        
        try:
            doc = Document(file_path)
            all_text = []
            current_char_position = 0
            page_mapping = []
            
            # Extract paragraphs
            paragraph_count = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraph_count += 1
                    para_start = current_char_position
                    para_end = current_char_position + len(text)
                    
                    page_mapping.append({
                        "paragraph_number": paragraph_count,
                        "start_char": para_start,
                        "end_char": para_end,
                        "text_length": len(text)
                    })
                    
                    all_text.append(text)
                    current_char_position = para_end + 1
            
            # Extract tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        all_text.append(row_text)
                        current_char_position += len(row_text) + 1
            
            full_text = "\n".join(all_text)
            
            # Metadata
            metadata = {
                "file_size": os.path.getsize(file_path),
                "file_name": Path(file_path).name,
                "file_type": "docx",
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "total_characters": len(full_text)
            }
            
            logger.info(f"Extracted {len(full_text)} characters from {paragraph_count} paragraphs and {table_count} tables")
            
            return ExtractionResult(
                text=full_text,
                metadata=metadata,
                page_mapping=page_mapping
            )
            
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract DOCX: {str(e)}")


class TXTExtractor(TextExtractor):
    """Extract text from plain text files with encoding detection."""
    
    def supports(self, file_extension: str) -> bool:
        """Check if file is a text file."""
        return file_extension.lower() in ['.txt', '.md', '.csv', '.log']
    
    def extract(self, file_path: str) -> ExtractionResult:
        """
        Extract text from plain text file with encoding detection.
        
        Args:
            file_path: Path to text file
            
        Returns:
            ExtractionResult with text and metadata
        """
        logger.info(f"Extracting text from TXT: {file_path}")
        
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
            
            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                text = f.read()
            
            # Create simple page mapping (no actual pages in text files)
            page_mapping = [{
                "section": "full_document",
                "start_char": 0,
                "end_char": len(text),
                "text_length": len(text)
            }]
            
            # Metadata
            metadata = {
                "file_size": os.path.getsize(file_path),
                "file_name": Path(file_path).name,
                "file_type": "txt",
                "encoding": encoding,
                "total_characters": len(text),
                "line_count": text.count('\n') + 1
            }
            
            logger.info(f"Extracted {len(text)} characters with {encoding} encoding")
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                page_mapping=page_mapping
            )
            
        except Exception as e:
            logger.error(f"Error extracting TXT {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract TXT: {str(e)}")


class ExtractorFactory:
    """Factory to get appropriate extractor based on file type."""
    
    def __init__(self):
        self.extractors = [
            PDFExtractor(),
            DOCXExtractor(),
            TXTExtractor()
        ]
    
    def get_extractor(self, file_path: str) -> TextExtractor:
        """
        Get appropriate extractor for file.
        
        Args:
            file_path: Path to file
            
        Returns:
            TextExtractor instance
            
        Raises:
            ValueError: If no extractor supports the file type
        """
        file_extension = Path(file_path).suffix
        
        for extractor in self.extractors:
            if extractor.supports(file_extension):
                return extractor
        
        raise ValueError(f"Unsupported file type: {file_extension}")
    
    def extract(self, file_path: str) -> ExtractionResult:
        """
        Extract text using appropriate extractor.
        
        Args:
            file_path: Path to file
            
        Returns:
            ExtractionResult
        """
        extractor = self.get_extractor(file_path)
        return extractor.extract(file_path)
